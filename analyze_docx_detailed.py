from docx import Document
import os

# DOCX 파일 분석
doc = Document('전기기능사 시험문제(2025-2023).docx')

print("=" * 80)
print("DOCX 파일 상세 분석 (그림 포함)")
print("=" * 80)

# 블록 단위 (단락, 테이블, 그림 등) 추출
def get_body_blocks(doc):
    """문단과 테이블, 그림을 포함한 모든 블록 추출"""
    blocks = []
    for para in doc.paragraphs:
        text = para.text.strip()
        # 그림 확인
        has_image = False
        for run in para.runs:
            if run._element.xpath('.//pic:pic'):
                has_image = True
        blocks.append({
            'type': 'paragraph',
            'text': text,
            'has_image': has_image
        })
    
    # 표도 추가
    for table in doc.tables:
        blocks.append({
            'type': 'table',
            'text': '[표]'
        })
    
    return blocks

blocks = get_body_blocks(doc)

print(f"\n총 {len(blocks)}개 블록 발견\n")
print("처음 30개 블록:")
print("-" * 80)

for i, block in enumerate(blocks[:30]):
    block_type = block['type']
    text = block['text'][:60] if block['text'] else '[빈 줄]'
    
    if block_type == 'paragraph':
        image_mark = " 🖼️ [그림 포함]" if block.get('has_image') else ""
        print(f"[{i:2d}] {block_type:10s} | {text}{image_mark}")
    else:
        print(f"[{i:2d}] {block_type:10s} | {text}")

# 전체 블록 분석
print("\n" + "=" * 80)
print("전체 블록 통계:")
print("-" * 80)

paragraph_count = sum(1 for b in blocks if b['type'] == 'paragraph')
image_count = sum(1 for b in blocks if b.get('has_image'))
table_count = sum(1 for b in blocks if b['type'] == 'table')

print(f"단락: {paragraph_count}개")
print(f"그림 포함 단락: {image_count}개")
print(f"표: {table_count}개")

# 블록을 텍스트로 저장
with open('docx_blocks_analysis.txt', 'w', encoding='utf-8') as f:
    for i, block in enumerate(blocks):
        f.write(f"[블록 {i}] {block['type']}\n")
        if block['text']:
            f.write(f"내용: {block['text']}\n")
        if block.get('has_image'):
            f.write(f"그림: 있음\n")
        f.write("\n")

print("\n✓ 분석 결과를 'docx_blocks_analysis.txt'에 저장했습니다")
